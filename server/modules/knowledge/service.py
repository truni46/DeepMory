# server/modules/knowledge/service.py
import asyncio
import hashlib
import os
import uuid
from pathlib import Path
from typing import Dict, List, Optional

from config.logger import logger
from modules.knowledge.repository import documentRepository
from modules.rag.ragService import ragService
from modules.llm.llmProvider import llmProvider
from common.prompts import DOCUMENT_SUMMARY_SYSTEM, documentSummaryUserPrompt
from modules.ocr.ocrProvider import ocrService, needsOcr

UPLOAD_DIR = Path(__file__).parent.parent.parent / "data" / "uploads"
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
OCR_DIR = Path(__file__).parent.parent.parent / "data" / "uploads" / "ocr"
OCR_DIR.mkdir(parents=True, exist_ok=True)

ALLOWED_TYPES = {".pdf", ".txt", ".md", ".docx", ".doc", ".xlsx", ".xls"}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 MB
SUMMARY_MAX_CHARS = 4000


def _computeHash(content: bytes) -> str:
    return hashlib.sha256(content).hexdigest()


def _extractPageCount(filePath: str) -> int:
    ext = os.path.splitext(filePath)[1].lower()
    try:
        if ext == ".pdf":
            import pypdf
            reader = pypdf.PdfReader(filePath)
            return len(reader.pages)
        elif ext in (".docx", ".doc"):
            import docx
            doc = docx.Document(filePath)
            paragraphs = [p for p in doc.paragraphs if p.text.strip()]
            return max(len(paragraphs) // 25, 1)
    except Exception as e:
        logger.warning(f"_extractPageCount failed for {filePath}: {e}")
    return 0


def _readTextContent(filePath: str, maxChars: int = SUMMARY_MAX_CHARS) -> str:
    ext = os.path.splitext(filePath)[1].lower()
    try:
        if ext == ".pdf":
            import pypdf
            reader = pypdf.PdfReader(filePath)
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif ext in (".docx", ".doc"):
            import docx
            doc = docx.Document(filePath)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        else:
            with open(filePath, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        return text[:maxChars]
    except Exception as e:
        logger.warning(f"_readTextContent failed for {filePath}: {e}")
        return ""


class DocumentService:

    async def uploadDocuments(
        self,
        userId: str,
        files: List,
        scope: str = "personal",
        ownerId: str = None,
        ownerType: str = "user",
    ) -> List[Dict]:
        results = []
        for fileObj in files:
            try:
                doc = await self._uploadOne(
                    userId, fileObj, scope, ownerId or userId, ownerType
                )
                results.append(doc)
            except Exception as e:
                logger.error(f"uploadDocuments failed for {fileObj.filename}: {e}")
                results.append({"error": str(e), "filename": fileObj.filename})
        return results

    async def _uploadOne(
        self,
        userId: str,
        fileObj,
        scope: str,
        ownerId: str,
        ownerType: str,
    ) -> Dict:
        filename = fileObj.filename
        fileExt = os.path.splitext(filename)[1].lower()
        if fileExt not in ALLOWED_TYPES:
            raise ValueError(f"File type '{fileExt}' not allowed")

        content = await fileObj.read()
        if len(content) > MAX_FILE_SIZE:
            raise ValueError("File exceeds 50 MB limit")

        contentHash = _computeHash(content)
        storedFilename = f"{uuid.uuid4().hex}_{filename}"
        filePath = UPLOAD_DIR / storedFilename

        with open(filePath, "wb") as f:
            f.write(content)

        record = await documentRepository.create(
            userId=userId,
            filename=filename,
            storedFilename=storedFilename,
            filePath=str(filePath),
            fileType=fileExt.lstrip("."),
            fileSize=len(content),
            contentHash=contentHash,
            scope=scope,
            ownerId=ownerId,
            ownerType=ownerType,
        )

        asyncio.create_task(
            self._processDocument(record["id"], str(filePath), ownerId, userId, filename)
        )
        return record

    async def _processDocument(
        self, documentId: str, filePath: str, ownerId: str, userId: str, filename: Optional[str] = None
    ) -> None:
        try:
            await documentRepository.updateEmbedding(documentId, "processing")

            indexPath = filePath
            if needsOcr(filePath):
                try:
                    await documentRepository.updateOcr(documentId, "processing", isScanned=True)
                    ocrPages = ocrService.ocrFile(filePath)
                    ocrFilePath = str(OCR_DIR / f"{documentId}_ocr.txt")
                    ocrService.saveOcrText(ocrPages, ocrFilePath)
                    await documentRepository.updateOcr(documentId, "completed", ocrFilePath=ocrFilePath)
                    indexPath = ocrFilePath
                    logger.info(f"OCR completed for {documentId}: {len(ocrPages)} pages")
                except Exception as e:
                    logger.error(f"OCR failed for {documentId}: {e}")
                    await documentRepository.updateOcr(documentId, "failed")

            chunkCount = await ragService.index(indexPath, ownerId, documentId, userId, filename=filename)
            pageCount = _extractPageCount(filePath)
            await documentRepository.updateEmbedding(
                documentId, "completed", chunkCount=chunkCount, pageCount=pageCount
            )
            logger.info(f"_processDocument completed for {documentId}")
            asyncio.create_task(self._generateSummary(documentId, filePath, indexPath))
        except Exception as e:
            logger.error(f"_processDocument failed for {documentId}: {e}")
            await documentRepository.updateEmbedding(
                documentId, "failed", errorMsg=str(e)
            )

    async def _generateSummary(self, documentId: str, filePath: str, indexPath: str = None) -> None:
        try:
            await documentRepository.updateSummary(documentId, "processing")
            text = _readTextContent(indexPath or filePath)
            if not text.strip():
                await documentRepository.updateSummary(documentId, "failed")
                return
            messages = [
                {"role": "system", "content": DOCUMENT_SUMMARY_SYSTEM},
                {"role": "user", "content": documentSummaryUserPrompt(text)},
            ]
            response = await llmProvider.generateResponse(messages, stream=False)
            await documentRepository.updateSummary(
                documentId, "completed", summary=response
            )
            logger.info(f"_generateSummary completed for {documentId}")
            doc = await documentRepository.getById(documentId)
            if doc:
                docUserId = doc.get("userId") or ""
                if docUserId:
                    asyncio.create_task(
                        ragService.upsertDocumentIndex(
                            userId=docUserId,
                            documentId=documentId,
                            filename=doc.get("filename", ""),
                            summary=response,
                        )
                    )
                else:
                    logger.warning(f"_generateSummary: skipping upsertDocumentIndex for {documentId} — missing userId")
        except Exception as e:
            logger.error(f"_generateSummary failed for {documentId}: {e}")
            await documentRepository.updateSummary(documentId, "failed")

    async def getDocuments(
        self, userId: str, scope: Optional[str] = None
    ) -> List[Dict]:
        return await documentRepository.getByUser(userId, scope)

    async def getDocument(self, documentId: str, userId: str) -> Optional[Dict]:
        doc = await documentRepository.getById(documentId)
        if not doc:
            logger.warning(f"getDocument: document {documentId} not found in DB")
            return None
        if str(doc.get("userId")) != str(userId):
            logger.warning(f"getDocument: userId mismatch for {documentId} — doc.userId={doc.get('userId')} vs request.userId={userId}")
            return None
        return doc

    async def getDocumentContext(self, documentIds: List[str], userId: str) -> str:
        parts = []
        for docId in documentIds:
            try:
                doc = await self.getDocument(docId, userId)
                if not doc:
                    logger.warning(f"getDocumentContext: document {docId} not found or unauthorized for user {userId}")
                    continue
                readPath = doc.get("ocrFilePath") or doc.get("filePath", "")
                if not readPath:
                    logger.warning(f"getDocumentContext: no readPath for document {docId}")
                    continue
                text = _readTextContent(readPath, maxChars=8000)
                if text.strip():
                    parts.append(f"--- Document: {doc.get('filename', docId)} ---\n{text}")
            except Exception as e:
                logger.error(f"getDocumentContext failed for document {docId}: {e}")
        return "\n\n".join(parts)

    async def searchDocumentContext(
        self, documentIds: List[str], userId: str, query: str
    ) -> tuple:
        try:
            firstDoc = await documentRepository.getById(documentIds[0])
            namespace = firstDoc.get("ownerId", userId) if firstDoc else userId

            results = await ragService.searchContextByDocumentIds(
                query=query, namespace=namespace, documentIds=documentIds, limit=5
            )
            maxScore = max((r.score for r in results), default=0.0)

            if results and maxScore >= 0.5:
                contextText = "\n\n".join(r.document.content for r in results)
                idToFilename: Dict[str, str] = {}
                for r in results:
                    did = r.document.metadata.get("documentId")
                    if did and did not in idToFilename:
                        d = await documentRepository.getById(did)
                        if d and d.get("filename"):
                            idToFilename[did] = d["filename"]
                sources = [
                    {
                        "filename": idToFilename.get(
                            r.document.metadata.get("documentId"),
                            r.document.metadata.get("filename"),
                        ),
                        "pageNumber": r.document.metadata.get("pageNumber"),
                        "documentId": r.document.metadata.get("documentId"),
                    }
                    for r in results
                ]
                return contextText, sources

            parts = []
            sources = []
            for docId in documentIds:
                try:
                    doc = await self.getDocument(docId, userId)
                    if not doc:
                        continue
                    readPath = doc.get("ocrFilePath") or doc["filePath"]
                    text = _readTextContent(readPath, maxChars=8000)
                    if text.strip():
                        parts.append(f"--- Document: {doc.get('filename', docId)} ---\n{text}")
                        sources.append({"filename": doc.get("filename"), "pageNumber": None, "documentId": docId})
                except Exception as e:
                    logger.error(f"searchDocumentContext fallback failed for {docId}: {e}")
            return "\n\n".join(parts), sources
        except Exception as e:
            logger.error(f"searchDocumentContext failed: {e}")
            return "", []

    async def deleteDocument(self, userId: str, documentId: str) -> bool:
        result = await documentRepository.delete(documentId, userId)
        if result is None:
            return False
        filePath, ownerId = result
        try:
            if filePath and os.path.exists(filePath):
                os.remove(filePath)
        except Exception as e:
            logger.error(f"deleteDocument file removal failed for {documentId}: {e}")
        ocrPath = str(OCR_DIR / f"{documentId}_ocr.txt")
        try:
            if os.path.exists(ocrPath):
                os.remove(ocrPath)
        except Exception as e:
            logger.error(f"deleteDocument OCR file removal failed for {documentId}: {e}")
        try:
            await ragService.deleteDocumentChunks(ownerId, documentId)
        except Exception as e:
            logger.error(f"deleteDocument RAG cleanup failed for {documentId}: {e}")
        # doc-index is keyed by userId (uploader), not ownerId (project/user namespace)
        try:
            await ragService.deleteDocumentIndex(userId, documentId)
        except Exception as e:
            logger.error(f"deleteDocument doc-index cleanup failed for {documentId}: {e}")
        return True


documentService = DocumentService()
