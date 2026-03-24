"""
Document loader — reads files from disk and returns List[Document].
Supported types: PDF, DOCX, TXT, MD, HTML.
"""
from __future__ import annotations

import os
from typing import Dict, List, Optional

from config.logger import logger
from modules.rag.repository import Document


class DocumentLoader:

    async def loadFromPath(
        self, filePath: str, metadata: Optional[Dict] = None
    ) -> List[Document]:
        ext = os.path.splitext(filePath)[1].lower()
        meta = metadata or {}
        try:
            if ext == ".pdf":
                return await self._loadPdf(filePath, meta)
            elif ext == ".docx":
                return await self._loadDocx(filePath, meta)
            elif ext in {".md", ".txt", ".html", ".htm"}:
                return await self._loadText(filePath, meta)
            else:
                logger.warning(f"Unsupported file type '{ext}', loading as plain text.")
                return await self._loadText(filePath, meta)
        except Exception as e:
            logger.error(f"Failed to load document '{filePath}': {e}")
            return []

    async def _loadPdf(self, filePath: str, metadata: Dict) -> List[Document]:
        import asyncio

        def _extract():
            import pypdf
            reader = pypdf.PdfReader(filePath)
            pages = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                if text.strip():
                    pages.append(Document(
                        content=text,
                        source=filePath,
                        metadata={"page": i + 1, **metadata},
                    ))
            return pages

        return await asyncio.get_event_loop().run_in_executor(None, _extract)

    async def _loadDocx(self, filePath: str, metadata: Dict) -> List[Document]:
        import asyncio

        def _extract():
            import docx
            doc = docx.Document(filePath)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return [Document(content=text, source=filePath, metadata=metadata)]

        return await asyncio.get_event_loop().run_in_executor(None, _extract)

    async def _loadText(self, filePath: str, metadata: Dict) -> List[Document]:
        import asyncio

        def _read():
            with open(filePath, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            return [Document(content=content, source=filePath, metadata=metadata)]

        return await asyncio.get_event_loop().run_in_executor(None, _read)


documentLoader = DocumentLoader()
