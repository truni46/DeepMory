"""
RAG public facade — supports 'simple' (direct Qdrant) and 'lightrag' providers.
Switch via RAG_PROVIDER env var (default: simple).
"""
from __future__ import annotations

import os
import uuid
from typing import Dict, List, Optional

from config.logger import logger
from modules.rag.repository import Document, SearchResult


def _readFile(filePath: str) -> str:
    """Read file content as text. Used by LightRagAdapter."""
    ext = os.path.splitext(filePath)[1].lower()
    try:
        if ext == ".pdf":
            import pypdf
            reader = pypdf.PdfReader(filePath)
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        elif ext == ".docx":
            import docx
            doc = docx.Document(filePath)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        else:
            with open(filePath, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
    except Exception as e:
        logger.error(f"_readFile failed for '{filePath}': {e}")
        return ""


def _toSearchResults(lightragResponse: str) -> List[SearchResult]:
    if not lightragResponse or not lightragResponse.strip():
        return []
    return [
        SearchResult(
            document=Document(
                id=str(uuid.uuid4()),
                content=lightragResponse.strip(),
                metadata={"source": "lightrag"},
            ),
            score=1.0,
        )
    ]


class LightRagAdapter:
    """Wraps lightragProvider — identical behavior to old RagService."""

    async def index(self, filePath: str, projectId: str, documentId: str, userId: str) -> int:
        try:
            from lightrag import QueryParam
            from modules.rag.lightragProvider import lightragProvider
            content = _readFile(filePath)
            if not content.strip():
                return 0
            instance = await lightragProvider.getInstance(f"project_{projectId}")
            await instance.ainsert(content, ids=[documentId], file_paths=[filePath])
            logger.info(f"LightRagAdapter: indexed document {documentId} for project {projectId}")
            return 1
        except Exception as e:
            logger.error(f"LightRagAdapter.index failed for document {documentId}: {e}")
            raise

    async def deleteDocumentChunks(self, projectId: str, documentId: str) -> None:
        try:
            from modules.rag.lightragProvider import lightragProvider
            instance = await lightragProvider.getInstance(f"project_{projectId}")
            await instance.adelete_by_doc_id(documentId)
        except Exception as e:
            logger.error(f"LightRagAdapter.deleteDocumentChunks failed for document {documentId}: {e}")
            raise

    async def searchContext(
        self, query: str, projectId: str, limit: int = 5, rerank: bool = False, mode: str = None
    ) -> List[SearchResult]:
        try:
            from lightrag import QueryParam
            from modules.rag.lightragProvider import lightragProvider
            instance = await lightragProvider.getInstance(f"project_{projectId}")
            queryMode = mode or os.getenv("LIGHTRAG_QUERY_MODE", "hybrid")
            result = await instance.aquery(query, param=QueryParam(mode=queryMode, top_k=limit))
            return _toSearchResults(result)
        except Exception as e:
            logger.warning(f"LightRagAdapter.searchContext failed for project {projectId}: {e}")
            return []

    async def searchContextByDocumentIds(
        self, query: str, namespace: str, documentIds: List[str], limit: int = 5
    ) -> List[SearchResult]:
        logger.warning("LightRagAdapter.searchContextByDocumentIds not supported — caller will use full-text fallback")
        return []

    async def upsertMemoryVector(
        self, userId: str, memoryId: str, content: str, metadata: Optional[Dict] = None
    ) -> None:
        try:
            from modules.rag.lightragProvider import lightragProvider
            instance = await lightragProvider.getInstance(f"user_{userId}")
            await instance.ainsert(content, ids=[memoryId])
        except Exception as e:
            logger.error(f"LightRagAdapter.upsertMemoryVector failed for user {userId}: {e}")

    async def searchMemoryVectors(
        self, userId: str, query: str, limit: int = 5
    ) -> List[SearchResult]:
        try:
            from lightrag import QueryParam
            from modules.rag.lightragProvider import lightragProvider
            instance = await lightragProvider.getInstance(f"user_{userId}")
            result = await instance.aquery(query, param=QueryParam(mode="hybrid", top_k=limit))
            return _toSearchResults(result)
        except Exception as e:
            logger.warning(f"LightRagAdapter.searchMemoryVectors failed for user {userId}: {e}")
            return []

    async def deleteMemoryVector(self, userId: str, memoryId: str) -> None:
        try:
            from modules.rag.lightragProvider import lightragProvider
            instance = await lightragProvider.getInstance(f"user_{userId}")
            await instance.adelete_by_doc_id(memoryId)
        except Exception as e:
            logger.error(f"LightRagAdapter.deleteMemoryVector failed memoryId={memoryId}: {e}")

    async def upsertDocumentIndex(
        self, userId: str, documentId: str, filename: str, summary: str
    ) -> None:
        logger.debug("LightRagAdapter.upsertDocumentIndex: not supported, skipping")

    async def searchDocumentIndex(
        self, userId: str, query: str, limit: int = 10, threshold: float = 0.6
    ) -> List[Dict]:
        return []

    async def deleteDocumentIndex(self, userId: str, documentId: str) -> None:
        logger.debug("LightRagAdapter.deleteDocumentIndex: not supported, skipping")


class RagService:
    """Public facade — delegates to SimpleRagProvider or LightRagAdapter based on RAG_PROVIDER env."""

    def __init__(self):
        providerName = os.getenv("RAG_PROVIDER", "simple").lower()
        if providerName == "lightrag":
            self._provider = LightRagAdapter()
            logger.info("RagService: using LightRAG provider")
        else:
            from modules.rag.simpleRagProvider import simpleRagProvider
            self._provider = simpleRagProvider
            logger.info("RagService: using Simple RAG provider (direct Qdrant + embeddings)")

    async def index(self, filePath: str, projectId: str, documentId: str, userId: str) -> int:
        return await self._provider.index(filePath, projectId, documentId, userId)

    async def deleteDocumentChunks(self, projectId: str, documentId: str) -> None:
        return await self._provider.deleteDocumentChunks(projectId, documentId)

    async def searchContext(
        self, query: str, projectId: str, limit: int = 5, rerank: bool = False, mode: str = None
    ) -> List[SearchResult]:
        return await self._provider.searchContext(query, projectId, limit=limit, rerank=rerank, mode=mode)

    async def searchContextByDocumentIds(
        self, query: str, namespace: str, documentIds: List[str], limit: int = 5
    ) -> List[SearchResult]:
        return await self._provider.searchContextByDocumentIds(
            query=query, namespace=namespace, documentIds=documentIds, limit=limit
        )

    async def upsertMemoryVector(
        self, userId: str, memoryId: str, content: str, metadata: Optional[Dict] = None
    ) -> None:
        return await self._provider.upsertMemoryVector(userId, memoryId, content, metadata)

    async def searchMemoryVectors(self, userId: str, query: str, limit: int = 5) -> List[SearchResult]:
        return await self._provider.searchMemoryVectors(userId, query, limit=limit)

    async def deleteMemoryVector(self, userId: str, memoryId: str) -> None:
        return await self._provider.deleteMemoryVector(userId, memoryId)

    async def upsertDocumentIndex(
        self, userId: str, documentId: str, filename: str, summary: str
    ) -> None:
        return await self._provider.upsertDocumentIndex(userId, documentId, filename, summary)

    async def searchDocumentIndex(
        self, userId: str, query: str, limit: int = 10, threshold: float = 0.6
    ) -> List[Dict]:
        return await self._provider.searchDocumentIndex(userId, query, limit=limit, threshold=threshold)

    async def deleteDocumentIndex(self, userId: str, documentId: str) -> None:
        return await self._provider.deleteDocumentIndex(userId, documentId)


ragService = RagService()
